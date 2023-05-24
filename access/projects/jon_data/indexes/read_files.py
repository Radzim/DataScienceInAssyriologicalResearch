import ast
import re
import sqlite3
import docx
import numpy as np
import pickle
import pandas as pd


def save_pkl(file, filename):
    with open('pickles/'+filename + '.pickle', 'wb') as handle:
        pickle.dump(file, handle, protocol=pickle.HIGHEST_PROTOCOL)


def load_pkl(filename):
    with open('pickles/'+filename + '.pickle', 'rb') as handle:
        return pickle.load(handle)


def get_name_text_dict(filename):
    try:
        return load_pkl('name_text_dict')
    except:
        doc = docx.Document(filename)
        full_text = {}
        for para in doc.paragraphs[87:10847]: # hardcoded to current doc!
            two_parts = para.text.split('\t')
            if len(two_parts) == 2:
                full_text[two_parts[0]] = two_parts[1].split('; ')
        save_pkl(full_text, 'name_text_dict')
        return full_text

#
# def get_text_name_dict(ntd):
#     try:
#         return load_pkl('text_name_dict')
#     except:
#         tnd = inverse_dict(ntd)
#         save_pkl(tnd, 'text_name_dict')
#         return tnd
#
#
# def inverse_dict(dictionary):
#     new_dict = {}
#     for key in dictionary.keys():
#         for sub_item in dictionary[key]:
#             if sub_item in new_dict.keys():
#                 new_dict[sub_item] += [key]
#             else:
#                 new_dict[sub_item] = [key]
#     return new_dict
#
#
# def make_text_name_matrix(t, n):
#     try:
#         return load_pkl('text_name_matrix')
#     except:
#         matrix = np.zeros([len(t), len(n)])
#         for name in n:
#             for text in name_text_dict[name]:
#                 matrix[t.index(text)][n.index(name)] = 1
#         save_pkl(matrix, 'text_name_matrix')
#         return matrix
#
#
# def make_text_text_matrix(t, tnm):
#     try:
#         return load_pkl('text_text_matrix')
#     except:
#         matrix = np.zeros([len(t), len(t)])
#         t2 = [tt for tt in t if sum(tnm[t.index(tt)]) > 5]
#         for text_a in t2:
#             for text_b in t2:
#                 count = np.dot(tnm[t.index(text_a)], tnm[t.index(text_b)])
#                 matrix[t.index(text_a)][t.index(text_b)] = count
#         save_pkl(matrix, 'text_text_matrix')
#         return matrix
#
#
# def n_largest_values(arr, t, n):
#     arr = arr.copy()
#     np.fill_diagonal(arr, 0)
#     ind = np.argpartition(arr.ravel(), -n*2)[-n*2:]
#     c = np.column_stack(np.unravel_index(ind, arr.shape))
#     c = [(t[cc[0]], t[cc[1]]) for cc in c if cc[0] > cc[1]]
#     return c

# text_name_dict = get_text_name_dict(name_text_dict)
# names, texts = list(name_text_dict.keys()), list(text_name_dict.keys())
# text_name_matrix = make_text_name_matrix(texts, names)
# text_text_matrix = make_text_text_matrix(texts, text_name_matrix)
#
# df = pd.DataFrame(n_largest_values(text_text_matrix, texts, 1000))
# df.columns = ['T1', 'T2']
# df['names'] = df.apply(lambda x: text_text_matrix[texts.index(x.T1)][texts.index(x.T2)], axis=1)
# df = df.sort_values(by='names', axis=0, ascending=False).reset_index(drop=True)
# df.to_csv('n_links.csv')

def create_csv(dict_):
    all_entries = []
    for name in dict_:
        entries = dict_[name]
        for entry in entries:
            all_entries.append([name, entry, '', '', ''])
    entries_csv = pd.DataFrame(all_entries, columns=['name', 'entry', 'text_number', 'line_numbers', 'additional_information'])
    entries_csv.to_csv('csvs/pn_entries.csv')
    return entries_csv

def partition_text(entry):
    references = ['12 N', '12 NT', '2 NT', '3N', '3 NT', '4 NT', 'A', 'AO', 'AfO', 'Akkadica', 'AS', 'Ashm.', 'AUAM', 'B.', 'Bab', 'BaF', 'BaM', 'BE', 'BM', 'Brinkman Dur-Karigalzu Catalog', 'CBS', 'CUNES', 'D', 'DK', 'EAH', 'FFA', 'HS', 'HSM', 'IM', 'L', 'MDP', 'MFA', 'MS', 'MSKH 1', 'MUN', 'N', 'Ni.', 'P', 'PBS', 'ROM', 'RS', 'Sb', 'TuM NF V', 'U.', 'UET', 'UM', 'VAT', 'WCMA', 'WHM', 'YBC']
    # entry = re.sub(" \(.*?\)",'', entry)
    tokens = entry.split(' ')
    results = []
    # 'text_number', 'line_numbers', 'additional_information'
    for i in range(len(tokens)-1):
        if tokens[i] in references:
            if len(tokens[i+1]) and tokens[i+1][0].isnumeric():
                string = ' '.join(tokens[i:])

                preends = [',', '(']
                prepositions = []
                for end in preends:
                    prepositions.append(string.find(end) % (len(string)+1))
                additional_information = ' '.join(tokens[:i]) + ' ' + string[min(prepositions):]
                string = string[:min(prepositions)]

                ends = [':', ';', ' rev', 'obv', ' i', ' ii', ' iii', ' iv', ' side', ' r', ' edge', '+', '=', 'smaller', 'larger']
                positions = []
                for end in ends:
                    positions.append(string.find(end) % (len(string)+1))
                text_number = string[:min(positions)]
                line_numbers = string[min(positions):min(prepositions)]
                results.append([text_number, line_numbers, additional_information])
    return results

def update_csv(_csv):
    rows = []
    for i, row in _csv.iterrows():
        x = partition_text(row.entry)
        if not len(x):
            rows.append(row)
        else:
            for j in range(len(x)):
                row2 = row.copy()
                row2.text_number = x[j][0]
                row2.line_numbers = x[j][1]
                row2.additional_information = x[j][2]
                rows.append(row2)
    _csv2 = pd.DataFrame(rows, columns=['name', 'entry', 'text_number', 'line_numbers', 'additional_information'])
    _csv2.to_csv('csvs/pn_entries.csv', index=False)
    return _csv

# name_text_dict = get_name_text_dict('PersonalNames.docx')
# name_text_csv = create_csv(name_text_dict)
# name_text_csv = update_csv(name_text_csv)

def get_place_text_dict(filename):
    try:
        return load_pkl('place_text_dict')
    except:
        doc = docx.Document(filename)
        full_text = {}
        for para in doc.paragraphs[0:]:  # hardcoded to current doc!
            two_parts = para.text.split(': ')
            if len(two_parts) == 2:
                full_text[two_parts[0]] = [partition_text(x) for x in two_parts[1].split('; ')]
        save_pkl(full_text, 'place_text_dict')
        return full_text

place_text_dict = get_place_text_dict('GeographicalNames.docx')

def place_text_df(pt_dict):
    big_df = []
    for i in pt_dict:
        pt_entry = pt_dict[i]
        new_pt = []
        for e in pt_entry:
            if len(e):
                new_pt += e
        pt_entry = new_pt
        if len(pt_entry):
            df = pd.DataFrame(pt_entry, columns=['text_number','line_numbers','additional_information'])
            df['place'] = i
            big_df.append(df[['place', 'text_number','line_numbers','additional_information']])
    big_df = pd.concat(big_df)
    big_df.to_csv('csvs/gn_entries.csv')
    return big_df

print(place_text_df(place_text_dict))

# def text_name_csv(nt_csv):
#     d = nt_csv.groupby('text_number')['name'].apply(list).to_dict()
#     a = list(d.keys())
#     b = list([set(dd) for dd in d.values()])
#     df = pd.DataFrame(zip(a, b), columns=['museum_no', 'personal_names'])
#     df.to_csv('csvs/tablets_to_pns.csv', index=False)
#
# def name_text_csv(nt_csv):
#     d = nt_csv.groupby('name')['text_number'].apply(list).to_dict()
#     a = list(d.keys())
#     b = list([set(filter(lambda x: x == x , dd)) for dd in d.values()])
#     df = pd.DataFrame(zip(a, b), columns=['personal_names', 'museum_no'])
#     df.to_csv('csvs/pns_to_tablets.csv', index=False)


# nt_csv = pd.read_csv('csvs/pn_entries.csv')
# text_name_csv(nt_csv)
# name_text_csv(nt_csv)

# df = pd.read_csv('csvs/tablets_to_pns.csv')
# df = df.set_index('museum_no')
# df.personal_names = df.personal_names.apply(ast.literal_eval)
#
# # save_pkl(df.personal_names.to_dict(), 'text_name_dict')
# print(df.personal_names.to_dict())
# print(load_pkl('name_text_dict'))
# set_ = df.loc['AO 2507']['personal_names']
# # df2 =  & set
# links_df = (set_ - (set_ - df['personal_names']))
# print(links_df.str.len().sort_values(ascending=False))
